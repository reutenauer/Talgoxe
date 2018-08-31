# -*- coding: utf-8 -*-
from __future__ import unicode_literals
import re
from tempfile import mkdtemp

import ezodf
from ezodf import newdoc, Heading, Paragraph, Span
from lxml import etree
import docx

from django.db import models
from django.core.exceptions import ObjectDoesNotExist

# FIXME Superscripts!

class Artikel(models.Model):
    lemma = models.CharField(max_length = 100)
    rang = models.SmallIntegerField()
    skapat = models.DateTimeField(auto_now_add = True)
    uppdaterat = models.DateTimeField(auto_now = True)
    segments = []

    def __str__(self):
        return self.lemma

    def get_spole(self, i):
        if i < self.spole_set.count():
            return self.spole_set.order_by('pos').all()[i] # Man kan missa vissa värden av pos, så .get(pos = i)funkar inte. Se t.ex. 1öden (id 3012683), spole saknas för pos = 2. AR 2018-06-03.
        else:
            return None

    def resolve_moments(self, segment):
        if segment.ism1():
            if len(self.moments['M2']) > 1:
                for m2 in range(len(self.moments['M2'])):
                    self.moments['M2'][m2].text = '%c' % (97 + m2)
                    self.moments['M2'][m2].display = True
            self.moments['M2'] = []
            self.moments['M1'].append(segment)
        elif segment.ism2():
            self.moments['M2'].append(segment)

    def append_segment(self, data):
        segment = Segment(data)
        self.resolve_moments(segment)
        self.new_segments.append(segment)

    def collect(self):
        self.new_segments = []
        i = 0
        state = 'INITIAL'
        self.moments = { 'M1': [], 'M2': [] }
        landskap = []
        while i < self.spole_set.count():
            currdat = self.get_spole(i)
            if state == 'GEOGRAFI':
                if currdat.isgeo():
                    landskap.append(Landskap(currdat.text))
                else:
                    sorted_landskap = sorted(landskap, key = Landskap.key)
                    for ls in sorted_landskap:
                       self.new_segments.append(Segment(geotype, ls.abbrev))
                    landskap = []
                    bits = re.split(u'¶', currdat.text) # För pilcrow i ”hårgård” och ”häringa”
                    if len(bits) == 1:
                        self.append_segment(currdat)
                    else:
                        maintype = currdat.typ
                        for bit in bits:
                            if bits.index(bit) > 0:
                                i += 1
                                if i < self.spole_set.count():
                                    self.new_segments.append(Segment(self.get_spole(i)))
                            if bit:
                                self.new_segments.append(Segment(maintype, bit))
                    state = 'INITIAL'
            else:
                if currdat.isgeo():
                    landskap = [Landskap(currdat.text)]
                    geotype = currdat.typ
                    state = 'GEOGRAFI'
                else:
                    bits = re.split(u'¶', currdat.text)
                    if len(bits) == 1:
                        self.append_segment(currdat)
                    else:
                        maintype = currdat.typ
                        for bit in bits:
                            if bits.index(bit) > 0:
                                i += 1
                                if i < self.spole_set.count():
                                    self.new_segments.append(Segment(self.get_spole(i)))
                            if bit:
                                self.new_segments.append(Segment(maintype, bit))
            i += 1
        if landskap: # För landskapsnamnet på slutet av ”häringa”, efter bugfixet ovan
            sorted_landskap = sorted(landskap, key = Landskap.key)
            for ls in sorted_landskap:
                self.new_segments.append(Segment(geotype, ls.abbrev))
        if len(self.moments['M1']) > 1:
            for m1 in range(len(self.moments['M1'])):
                self.moments['M1'][m1].text = '%d' % (m1 + 1)
                self.moments['M1'][m1].display = True
        if len(self.moments['M2']) > 1:
            for m2 in range(len(self.moments['M2'])):
                self.moments['M2'][m2].text = '%c' % (97 + m2)
                self.moments['M2'][m2].display = True
        self.moments = { 'M1': [], 'M2': [] }

    def process(self, outfile):
        self.collect()
        outfile.write("\\hskip-0.5em")
        if self.rang > 0:
            outfile.write('\lohi[left]{}{\SDL:SO{%d}}' % self.rang) # TODO Real superscript
        outfile.write("\\SDL:SO{%s}" % self.lemma)
        setspace = True
        for segment in self.new_segments: # TODO Handle moments!  segment.ismoment and segment.display
            if setspace and not segment.isrightdelim():
                outfile.write(' ') # FIXME But not if previous segment is left delim!
            if segment.isleftdelim():
                setspace = False
            else:
                setspace = True
            type = segment.type.__str__()
            text = segment.format().replace(u'\\', '\\textbackslash ').replace('~', '{\\char"7E}')
            outfile.write(('\\SDL:%s{%s}' % (type, text)))

    @staticmethod
    def add_style(opendoc, type, xmlchunk):
        opendoc.inject_style("""
            <style:style style:name="%s" style:family="text">
                <style:text-properties %s />
            </style:style>
            """ % (type, xmlchunk))

    @staticmethod
    def start_odf(tempfilename):
        odt = ezodf.newdoc(doctype = 'odt', filename = tempfilename)
        Artikel.add_style(odt, 'SO', 'fo:font-weight="bold"')
        Artikel.add_style(odt, 'OK', 'fo:font-size="9pt"')
        Artikel.add_style(odt, 'G', 'fo:font-size="9pt"')
        Artikel.add_style(odt, 'DSP', 'fo:font-style="italic"')
        Artikel.add_style(odt, 'TIP', 'fo:font-size="9pt"')
        # IP
        Artikel.add_style(odt, 'M1', 'fo:font-weight="bold"')
        Artikel.add_style(odt, 'M2', 'fo:font-weight="bold"')
        # VH, HH, VR, HR
        Artikel.add_style(odt, 'REF', 'fo:font-weight="bold"')
        Artikel.add_style(odt, 'FO', 'fo:font-style="italic"')
        Artikel.add_style(odt, 'TIK', 'fo:font-style="italic" fo:font-size="9pt"')
        Artikel.add_style(odt, 'FLV', 'fo:font-weight="bold" fo:font-size="9pt"')
        # ÖVP. Se nedan!
        # BE, ÖV
        # ÄV, ÄVK se nedan
        # FOT
        Artikel.add_style(odt, 'GT', 'fo:font-size="9pt"')
        Artikel.add_style(odt, 'SOV', 'fo:font-weight="bold"')
        # TI
        # HV, INT
        Artikel.add_style(odt, 'OKT', 'fo:font-size="9pt"')
        # VS
        Artikel.add_style(odt, 'GÖ', 'fo:font-size="9pt"')
        # GP, UST
        Artikel.add_style(odt, 'US', 'fo:font-style="italic"')
        # GÖP, GTP, NYR, VB
        Artikel.add_style(odt, 'OG', 'style:text-line-through-style="solid"')
        Artikel.add_style(odt, 'SP', 'fo:font-style="italic"')

        return odt

    def process_odf(self, odt):
        paragraph = self.generate_content()
        odt.body += paragraph

    @staticmethod
    def stop_odf(odt):
        odt.save()

    def generate_content(self):
        paragraph = Paragraph()
        paragraph += Span(self.lemma, style_name = 'SO') # TODO Homografnumrering!
        self.collect()
        spacebefore = True
        for segment in self.new_segments:
            type = segment.type.__str__()
            if not type == 'KO':
                if spacebefore and not segment.isrightdelim():
                    paragraph += Span(' ')
                paragraph += Span(segment.format(), style_name = type)
                if segment.isleftdelim():
                    spacebefore = False
                else:
                    spacebefore = True
        return paragraph

    @staticmethod
    def add_docx_styles(document): # TODO Keyword arguments?
        Artikel.add_docx_style(document, 'SO', False, True, 12)
        Artikel.add_docx_style(document, 'OK', False, False, 9)
        Artikel.add_docx_style(document, 'G', False, False, 9)
        Artikel.add_docx_style(document, 'DSP', True, False, 12)
        Artikel.add_docx_style(document, 'TIP', False, False, 9)
        Artikel.add_docx_style(document, 'IP', False, False, 12)
        Artikel.add_docx_style(document, 'M1', False, True, 12)
        Artikel.add_docx_style(document, 'M2', False, True, 12)
        Artikel.add_docx_style(document, 'VH', False, False, 12)
        Artikel.add_docx_style(document, 'HH', False, False, 12)
        Artikel.add_docx_style(document, 'VR', False, False, 12)
        Artikel.add_docx_style(document, 'HR', False, False, 12)
        Artikel.add_docx_style(document, 'REF', False, True, 12)
        Artikel.add_docx_style(document, 'FO', True, False, 12)
        Artikel.add_docx_style(document, 'TIK', True, False, 9)
        Artikel.add_docx_style(document, 'FLV', False, True, 9)
        Artikel.add_docx_style(document, 'ÖVP', False, False, 12)
        Artikel.add_docx_style(document, 'BE', False, False, 12)
        Artikel.add_docx_style(document, 'ÖV', False, False, 12)
        Artikel.add_docx_style(document, 'ÄV', False, False, 12) # FIXME Skriv “även” :-)
        Artikel.add_docx_style(document, 'ÄVK', True, False, 12) # FIXME Skriv även även här ;-)
        Artikel.add_docx_style(document, 'FOT', True, False, 12)
        Artikel.add_docx_style(document, 'GT', False, False, 9)
        Artikel.add_docx_style(document, 'SOV', False, True, 12)
        for style in ('TI', 'HV', 'INT'):
            Artikel.add_docx_style(document, style)
        Artikel.add_docx_style(document, 'OKT', False, False, 9)
        Artikel.add_docx_style(document, 'VS')
        Artikel.add_docx_style(document, 'GÖ', False, False, 9)
        Artikel.add_docx_style(document, 'GP')
        Artikel.add_docx_style(document, 'UST', True, False, 12)
        Artikel.add_docx_style(document, 'US', True, False, 12)
        for style in ('GÖP', 'GTP', 'NYR', 'VB'):
            Artikel.add_docx_style(document, style)
        OG = document.styles.add_style('OG', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        OG.font.strike = True
        Artikel.add_docx_style(document, 'SP', True, False, 12)
        Artikel.add_docx_style(document, 'M0', False, True, 18)
        Artikel.add_docx_style(document, 'M3', True, False, 6)

    @staticmethod
    def add_docx_style(document, type, italic = False, bold = False, size = 12):
        style = document.styles.add_style(type, docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        style.base_style = document.styles['Default Paragraph Font']
        if italic:
            style.font.italic = True
        if bold:
            style.font.bold = True
        style.font.size = docx.shared.Pt(size)

    def process_docx(self, filename):
        document = docx.Document()
        Artikel.add_docx_styles(document)
        self.generate_docx_paragraph(document)
        document.save(filename)

    def generate_docx_paragraph(self, document):
        self.collect()
        paragraph = document.add_paragraph()
        if self.rang > 0:
            paragraph.add_run(str(self.rang), style = 'SO').font.superscript = True
        paragraph.add_run(self.lemma, style = 'SO')
        spacebefore = True
        for segment in self.new_segments:
            type = segment.type.__str__()
            if not type == 'KO':
                if spacebefore and not segment.isrightdelim():
                    paragraph.add_run(' ', style = document.styles[type])
                if type == 'ÖVP': # TODO Något bättre
                    run = '(' + segment.format() + ')'
                else:
                    run = segment.format()
                paragraph.add_run(run, style = document.styles[type])
                if segment.isleftdelim():
                    spacebefore = False
                else:
                    spacebefore = True

    def serialise(self):
        paragraph = self.generate_content()
        return etree.tostring(paragraph.xmlnode)

    def update(self, post_data):
        order = post_data['order'].split(',')
        stickord = post_data['stickord']
        if self.lemma != stickord:
            self.lemma = stickord
            self.save()
        d = [[post_data['type-' + key].strip(), post_data['value-' + key].strip()] for key in order]
        gtype = Typ.objects.get(kod = 'g')
        for i in range(len(d)):
            bit = d[i]
            try:
                type = Typ.objects.get(kod = bit[0])
            except ObjectDoesNotExist: # FIXME Do something useful!
# TODO >>> Type.objects.create(abbrev = 'OG', name = 'Ogiltig', id = 63)
                type = Typ.objects.get(kod = 'OG')
            text = bit[1]
            if type == gtype and text.title() in Landskap.short_abbrev.keys():
              text = Landskap.short_abbrev[text.title()]
            data = self.get_spole(i)
            if data:
                data.typ = type
                data.text = text
                data.save()
            else:
                Spole.objects.create(artikel = self, typ = type, pos = i, text = text)
        self.spole_set.filter(pos__gte = len(d)).delete()

class Segment(): # Fjäder!
    def __init__(self, type, text = None):
        self.display = None
        if text != None:
            self.type = type
            self.text = text
        else: # type is actually a Data object
            self.type = type.typ
            self.text = type.text

    def __str__(self):
        return self.type.__str__() + ' ' + self.text

    def isgeo(self):
        return self.type.isgeo()

    def isleftdelim(self):
        if type(self.type) == 'unicode':
            return self.type in ['vh', 'vr']
        else:
            return self.type.isleftdelim()

    def isrightdelim(self):
        if type(self.type) == 'unicode' or str(type(self.type)) == "<type 'unicode'>":
            return self.type in ['hh', 'hr', 'ip', 'ko'] # KO is convenient
        else:
            return self.type.isrightdelim() or re.match('^,', self.text) # TODO Complete

    def ismoment(self):
        return self.ism1() or self.ism2()

    def ism1(self):
        return self.type.ism1()

    def ism2(self):
        return self.type.ism2()

    def output(self, outfile, next):
        outfile.write(('\SDL:%s{' % self.type).encode('UTF-8'))
        outfile.write(self.text.replace(u'\\', '\\backslash ').encode('UTF-8'))
        outfile.write('}')
        if not next.isrightdelim():
            outfile.write(' ')

    # TODO Method hyphornot()

    def format(self):
        if self.type.__str__().upper() == 'VH':
            return '['
        elif self.type.__str__().upper() == 'HH':
            return ']'
        elif self.type.__str__().upper() == 'VR':
            return '('
        elif self.type.__str__().upper() == 'HR':
            return ')'
        elif self.type.__str__().upper() == 'ÄV':
            return 'äv.'
        else:
            return self.text.strip()

class Typ(models.Model):
    kod = models.CharField(max_length = 5)
    namn = models.CharField(max_length = 30)
    skapat = models.DateTimeField(auto_now_add = True)
    uppdaterat = models.DateTimeField(auto_now = True)

    def __str__(self):
        return self.kod.upper()

    def isgeo(self):
        return self.kod == 'g'

    def isleftdelim(self):
        return self.kod in ['vh', 'vr']

    def isrightdelim(self):
        return self.kod in ['hh', 'hr', 'ip', 'ko']

    def ismoment(self):
        return self.ism1() or self.ism2()

    def ism1(self):
        return self.kod == 'm1'

    def ism2(self):
        return self.kod == 'm2'

    def format(self):
        out = self.__str__()
        out += (4 - len(out)) * '\xa0'
        return out

class Spole(models.Model):
    text = models.CharField(max_length = 2000)
    pos = models.SmallIntegerField()
    artikel = models.ForeignKey(Artikel)
    typ = models.ForeignKey(Typ)
    skapat = models.DateTimeField(auto_now_add = True)
    uppdaterat = models.DateTimeField(auto_now = True)

    def __str__(self):
        return self.typ.__str__() + ' ' + self.text

    def webstyle(self):
        self.webstyles[self.typ.__str__()]

    def printstyle(self):
        self.printstyles[self.typ.__str__()]

    def isgeo(self):
        return self.typ.isgeo()

class Landskap():
    ordning = [
        u'Skåne', u'Blek', u'Öland', u'Smål', u'Hall', u'Västg', u'Boh', u'Dalsl', u'Gotl', u'Östg', # 0-9
        u'Götal', # 10
        u'Sörml', u'Närke', u'Värml', u'Uppl', u'Västm', u'Dal', # 11 - 16
        u'Sveal', # 17
        u'Gästr', u'Häls', u'Härj', u'Med', u'Jämtl', u'Ång', u'Västb', u'Lappl', u'Norrb', # 18 - 26
        u'Norrl', # 27
    ]

    # TODO: Something more object-oriented (methods short, long, rank)
    short_abbrev = {
      u'Sk' : u'Skåne',
      u'Bl' : u'Blek',
      u'Öl' : u'Öland',
      u'Sm' : u'Smål',
      u'Ha' : u'Hall',
      u'Vg' : u'Västg',
      u'Bo' : u'Boh',
      u'Dsl': u'Dalsl',
      u'Gl' : u'Gotl',
      u'Ög' : u'Östg',
      u'Götal' : u'Götal',
      u'Sdm': u'Sörml',
      u'Nk' : u'Närke',
      u'Vrm': u'Värml',
      u'Ul' : u'Uppl',
      u'Vstm' : u'Västm',
      u'Dal': u'Dal',
      u'Sveal' : u'Sveal',
      u'Gst': u'Gästr',
      u'Hsl': u'Häls',
      u'Hrj': u'Härj',
      u'Mp' : u'Med',
      u'Jl' : u'Jämtl',
      u'Åm' : u'Ång',
      u'Vb' : u'Västb',
      u'Lpl': u'Lappl',
      u'Nb' : u'Norrb',
      u'Norrl' : u'Norrl',
    }

    def cmp(self, other):
        if self.abbrev in self.ordning and other.abbrev in self.ordning:
            return cmp(self.ordning.index(self.abbrev), self.ordning.index(other.abbrev))
        else:
            return 0

    @staticmethod
    def key(self):
        if self.abbrev in self.ordning:
            return self.ordning.index(self.abbrev)
        else:
            return -1 # Så det är lättare att se dem

    def __init__(self, abbrev):
        self.abbrev = abbrev.capitalize()

    def __str__(self):
        return self.abbrev
